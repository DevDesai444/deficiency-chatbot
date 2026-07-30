"""Tests for the DOCX parse path (Plan 04 / INGEST-02, INGEST-05, D-17, D-20, D-31).

Endpoint-free: uses the COMMITTED fixtures (merged_cells.docx from Plan 01, mini_spec.docx from
Phase 0) plus tmp-built edge-case docs. No live OCR / LLM.
"""
from __future__ import annotations

import docx
import pytest

import parse.docx as docxmod
from parse.docx import extract_docx
from parse.section_splitter import split_document

_MERGED = "tests/ingest/fixtures/merged_cells.docx"
_MINISPEC = "src/evals/dataset/docs/mini_spec.docx"
_PDF_KEYS = {"filename", "page_count", "toc", "pages"}
_PDF_PAGE_KEYS = {
    "page_number", "page_label", "width", "height", "rotation",
    "source", "is_scanned", "blocks", "tables", "figures",
}


def test_dict_shape_matches_extract_pdf():
    d = extract_docx(_MINISPEC)
    assert set(d) == _PDF_KEYS, set(d)
    assert d["page_count"] is None
    page = d["pages"][0]
    assert set(page) == _PDF_PAGE_KEYS, set(page)
    assert page["page_number"] is None and page["source"] == "python-docx"
    # geometry nulled on blocks (D-20)
    for b in page["blocks"]:
        assert b["bbox"] is None and b["page"] is None


def test_merged_cells_resolve_to_origin():
    d = extract_docx(_MERGED)
    tables = d["pages"][0]["tables"]
    assert tables, "no table parsed"
    t = tables[0]
    # covered coords map to their origin; origin coords are NOT keys (D-31)
    assert t["merged_origins"] == {"0,1": [0, 0], "2,2": [1, 2]}, t["merged_origins"]
    # table_id encodes the block-level reading-order position (deterministic, unique per doc)
    assert t["table_id"].startswith("docx-t")


def test_minispec_table_reconstructs_planted_deficiencies():
    d = extract_docx(_MINISPEC)
    tables = d["pages"][0]["tables"]
    assert tables
    t = tables[0]
    assert t["headers"] == ["Impurity", "Result (%)", "Limit"], t["headers"]
    flat = [c for row in t["rows"] for c in row]
    # the 3 planted deficiencies live in these rows (Total impurities / Maximum)
    assert "Total impurities" in flat and "Maximum" in flat, flat


def test_table_edge_cases(tmp_path, monkeypatch):
    # (a) borderless table (no "Table Grid" style) still enumerates its rows
    p = tmp_path / "borderless.docx"
    doc = docx.Document()
    tbl = doc.add_table(rows=2, cols=2)  # no style set -> borderless
    tbl.cell(0, 0).text = "H1"; tbl.cell(0, 1).text = "H2"
    tbl.cell(1, 0).text = "a"; tbl.cell(1, 1).text = "b"
    doc.save(str(p))
    d = extract_docx(p)
    t = d["pages"][0]["tables"][0]
    assert t["headers"] == ["H1", "H2"] and t["rows"] == [["a", "b"]]

    # (b) a nested table (cell.add_table) is reachable without crashing
    p2 = tmp_path / "nested.docx"
    doc2 = docx.Document()
    outer = doc2.add_table(rows=1, cols=1)
    outer.cell(0, 0).add_table(rows=1, cols=1)
    doc2.save(str(p2))
    d2 = extract_docx(p2)  # must not raise
    assert d2["pages"][0]["tables"], "outer table lost"

    # (c) complex/unreadable merge -> typed ParseFailed marker, never a crash (D-17, Pitfall 5)
    def _boom(*a, **k):
        raise IndexError("list index out of range")  # the documented python-docx complex-merge bug
    monkeypatch.setattr(docxmod, "_docx_table_to_extracted", _boom)
    d3 = extract_docx(_MINISPEC)
    marker = d3["pages"][0]["tables"][0]
    assert "_parse_failed" in marker
    assert marker["_parse_failed"]["layer"] == "parse.docx"
    assert "complex merge unreadable" in marker["_parse_failed"]["reason"]


def test_split_document_consumes_docx_dict():
    sections = split_document(extract_docx(_MINISPEC))
    assert len(sections) >= 1
