"""Deterministic generator for the merged-cell DOCX fixture.

Writes `tests/ingest/fixtures/merged_cells.docx`: a 3x3 table exercising BOTH
merge types python-docx heals to a single `_tc` element (RESEARCH.md Pattern 3,
the empirical basis for D-31 "every covered coordinate resolves to one origin"):

    col:      0                1         2
    row 0 | "Merged Header AB" (span) | "R0C2"
    row 1 | "R1C0"          | "R1C1"   | "Spanning Limit" (span
    row 2 | "R2C0"          | "R2C1"   |  continues)

  * (0,0)-(0,1) horizontally merged  -> origin (0,0), text "Merged Header AB"
  * (1,2)-(2,2) vertically merged    -> origin (1,2), text "Spanning Limit"

This is the ONLY merged-cell DOCX in the repo (`src/evals/dataset/docs/mini_spec.docx`
has no merges), and it is what SC2's merged-cell fidelity test gates on.

Idempotent: re-running overwrites the file with the same deterministic content
(no random or time-based values written by us).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocumentT

OUTPUT_PATH = Path(__file__).parent / "merged_cells.docx"

_HEADING_TEXT = "Merged-Cell Fixture - Drug Substance Specification"
_NARRATIVE_TEXT = (
    "Fixture table exercising a horizontal (gridSpan) and a vertical (vMerge) "
    "merge so DOCX merged-cell resolution can be tested against python-docx's "
    "healed _tc grid model."
)


def build_document() -> DocumentT:
    """Construct the 3x3 merged-cell Document in memory."""
    document = Document()
    document.add_heading(_HEADING_TEXT, level=1)
    document.add_paragraph(_NARRATIVE_TEXT)

    table = document.add_table(rows=3, cols=3)
    table.style = "Table Grid"

    # Merge FIRST, then assign text to the origin, so python-docx does not
    # concatenate the pre-existing cell texts into the merged cell.
    top = table.cell(0, 0).merge(table.cell(0, 1))
    top.text = "Merged Header AB"          # origin (0,0)

    right = table.cell(1, 2).merge(table.cell(2, 2))
    right.text = "Spanning Limit"          # origin (1,2)

    # Distinct texts for every remaining coordinate so each is distinguishable.
    table.cell(0, 2).text = "R0C2"
    table.cell(1, 0).text = "R1C0"
    table.cell(1, 1).text = "R1C1"
    table.cell(2, 0).text = "R2C0"
    table.cell(2, 1).text = "R2C1"

    return document


def main() -> None:
    document = build_document()
    document.save(str(OUTPUT_PATH))
    print(f"wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
