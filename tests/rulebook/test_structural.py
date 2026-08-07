"""Tests for rulebook.structural — RECALL-02 deterministic structural inconsistency pass.

Ruling 5 — REAL fixture table shape:
Tests construct corpus cache entries matching the ACTUAL doc_b.docx two-column table
layout ("Impurity | % w/w") used in the synthetic fixture. This is NOT a simplified
one-column table — it uses the same paired label-column / value-column structure that
the real fixture has:
  Col 0 = Impurity names ("Compound A", "Compound B", "Total Impurities")
  Col 1 = Values ("0.10%", "0.18%", "0.12%")

The aggregate label "Total Impurities" is in col 0 (label column).
The claim value "0.12%" is in col 1 of the same row (value column).
Basis values are col 1 cells for all non-aggregate rows: "0.10%" and "0.18%".
True SUM = 0.10 + 0.18 = 0.28 → mismatch with claimed 0.12 → VIOLATION.
"""
from __future__ import annotations

import pytest
from ingest.anchors import mint_span
from ingest.manifest import CoverageManifest, DocEntry
from tests.tools.conftest import build_corpus_index
from tools.ledger import RetrievalLedger

# Guard: skip all tests if rulebook.structural not yet implemented
structural = pytest.importorskip(
    "rulebook.structural",
    reason="rulebook.structural not yet implemented",
)
detect_structural_inconsistencies = structural.detect_structural_inconsistencies
compare_values = structural.compare_values
_stated_precision = structural._stated_precision


# ---------------------------------------------------------------------------
# Fixture helpers
# ---------------------------------------------------------------------------

def _block(text, page=1, order=0):
    """Minimal block dict (same shape as extract_pdf emits)."""
    return {"text": text, "page": page, "reading_order": order, "lines": []}


def _two_col_table(table_id: str, headers, data_rows, page=1):
    """Build an ExtractedTable dict (grid kind) with a given table_id.

    Ruling 5: two-column layout matching the REAL doc_b.docx fixture shape.
    headers = [col0_header, col1_header]
    data_rows = [[col0_val, col1_val], ...]
    """
    return {
        "kind": "grid",
        "title": "",
        "headers": headers,
        "rows": data_rows,
        "pairs": [],
        "page": page,
        "bbox": (72, 100, 540, 400),
        "n_cols": 2,
        "n_rows": len(data_rows) + 1,  # +1 for header row
        "source_pages": [page],
        "continues_from": False,
        "continues_to": False,
        "table_id": table_id,
        "merged_origins": {},
    }


def _corpus_with_table(tmp_path, doc_id: str, headers, data_rows, table_id="t0"):
    """Build a CorpusIndex with a doc that has one addressable two-column table.

    Uses build_corpus_index from tests.tools.conftest which routes through the REAL
    serialize_document -> normalize -> build_table_index pipeline so cached_entry()
    returns a genuine cache entry shape (not a hand-rolled fake).
    """
    table = _two_col_table(table_id, headers, data_rows)
    blocks = [_block("Test document with impurity profile table.")]
    corpus = build_corpus_index(tmp_path, doc_id, blocks, tables=[table])
    return corpus


def _manifest_for(corpus, doc_id: str) -> CoverageManifest:
    """Return the corpus's manifest — build_corpus_index already sets tables='addressable'
    when the table_index is non-empty."""
    return corpus.manifest


# ---------------------------------------------------------------------------
# Test 1: Two-column aggregate violation detected (Ruling 5 — REAL fixture shape)
# ---------------------------------------------------------------------------

def test_aggregate_violation_detected_two_column_table(tmp_path):
    """Labeled-aggregate SUM violation: Total != sum of components.

    Ruling 5: REAL two-column table shape matching doc_b.docx Table 1 (X2b violation):
      Col 0 = Impurity names (label column)
      Col 1 = % w/w values (value column)
      Row 0: ["Impurity", "% w/w"]        <- headers
      Row 1: ["Compound A", "0.10%"]      <- basis
      Row 2: ["Compound B", "0.18%"]      <- basis
      Row 3: ["Total Impurities", "0.12%"] <- aggregate label (SUM expected 0.28%)

    The algorithm must:
    1. Find value column (col 1 has most numeric text)
    2. Find aggregate label at (row 3, col 0): "Total Impurities" in AGGREGATE_LEXICON
    3. Claim cell = (row 3, col 1) = "0.12%" (same row, value column)
    4. Basis cells = (row 1, col 1) = "0.10%", (row 2, col 1) = "0.18%"
    5. SUM(0.10, 0.18) = 0.28 != 0.12 -> VIOLATION
    """
    # Build corpus with real two-column table matching fixture doc_b.docx Table 1
    corpus = _corpus_with_table(
        tmp_path,
        doc_id="doc-b-table1",
        headers=["Impurity", "% w/w"],
        data_rows=[
            ["Compound A", "0.10%"],
            ["Compound B", "0.18%"],
            ["Total Impurities", "0.12%"],  # aggregate row: SUM mismatch
        ],
        table_id="t0",
    )
    manifest = _manifest_for(corpus, "doc-b-table1")
    ledger = RetrievalLedger()

    faults = detect_structural_inconsistencies(corpus, manifest, ledger)

    assert len(faults) >= 1, (
        "Ruling 5: detect_structural_inconsistencies must return at least one Fault "
        "for the two-column table where Total Impurities (0.12%) != SUM(0.10%, 0.18%)"
    )

    # Verify the fault has the right shape
    fault = faults[0]
    assert fault.leg_tag == "STRUCTURAL", f"leg_tag must be STRUCTURAL, got {fault.leg_tag}"
    assert fault.structural_anchor is not None, "structural_anchor must be set"
    anchor = fault.structural_anchor
    assert anchor.relation in ("SUM", "MAX"), (
        f"relation must be SUM or MAX for Total Impurities, got {anchor.relation}"
    )
    # The claim value must contain "0.12" (the stated but wrong total)
    assert "0.12" in anchor.actual_value, (
        f"actual_value must contain '0.12', got {anchor.actual_value!r}"
    )
    # The basis must have at least one span referencing "0.18" (Compound B)
    assert len(anchor.basis_span_ids) >= 1, "basis_span_ids must have at least one span"

    # Verify tier (Ruling 2)
    assert fault.tier.value == "verified", f"tier must be 'verified', got {fault.tier}"
    assert fault.evidence_class.value == "code_verified", (
        f"evidence_class must be 'code_verified', got {fault.evidence_class}"
    )


# ---------------------------------------------------------------------------
# Test 2: D-STR4 precision-derived comparison — compliant case
# ---------------------------------------------------------------------------

def test_precision_derived_complies(tmp_path):
    """D-STR4: NMT 0.10 vs 0.104 -> no violation (precision-derived rounding).

    compare_values('0.104', '0.10', 'LEQ') == False because:
    - prec = min(stated_precision('0.104'), stated_precision('0.10')) = min(3, 2) = 2
    - round(0.104, 2) = 0.10, round(0.10, 2) = 0.10
    - 0.10 <= 0.10 -> NOT a violation

    This is the critical D-STR4 correctness test: naive exact comparison would give
    0.104 > 0.10 = True (false positive), but precision-derived comparison gives False.
    """
    # Direct comparator test
    result = compare_values("0.104", "0.10", "LEQ")
    assert result == False, (  # noqa: E712 — must be exactly False, not None
        f"D-STR4: compare_values('0.104', '0.10', 'LEQ') must return False (complies), got {result}"
    )

    # Also verify the violation case
    result_violation = compare_values("0.15", "0.10", "LEQ")
    assert result_violation == True, (  # noqa: E712
        f"compare_values('0.15', '0.10', 'LEQ') must return True (violates), got {result_violation}"
    )

    # Abstain on unparseable
    result_abstain = compare_values("abc", "0.10", "LEQ")
    assert result_abstain is None, (
        f"compare_values('abc', '0.10', 'LEQ') must return None (abstain), got {result_abstain}"
    )

    # Verify a table with a compliant sum does NOT produce a fault
    # Table: Total = 0.28%, Compound A = 0.10%, Compound B = 0.18% -> SUM = 0.28 = OK
    corpus = _corpus_with_table(
        tmp_path,
        doc_id="doc-compliant",
        headers=["Impurity", "% w/w"],
        data_rows=[
            ["Compound A", "0.10%"],
            ["Compound B", "0.18%"],
            ["Total Impurities", "0.28%"],  # correct: SUM = 0.28
        ],
    )
    manifest = _manifest_for(corpus, "doc-compliant")
    ledger = RetrievalLedger()

    faults = detect_structural_inconsistencies(corpus, manifest, ledger)
    assert faults == [], (
        f"D-STR4: compliant table (Total 0.28% = SUM 0.10+0.18=0.28%) must produce NO faults, "
        f"got {len(faults)} fault(s): {[f.title for f in faults]}"
    )


# ---------------------------------------------------------------------------
# Test 3: Unavailable table tier skipped (D-STR5)
# ---------------------------------------------------------------------------

def test_unavailable_table_tier_skipped(tmp_path):
    """D-STR5: doc with empty table_index -> zero faults, no exception.

    The manifest entry for the doc has tables='unavailable'. The detector must
    log and skip without crashing, returning an empty list.
    """
    # Build corpus WITHOUT any tables (build_corpus_index sets tables='unavailable'
    # when table_index is empty)
    corpus = build_corpus_index(
        tmp_path,
        "doc-no-tables",
        blocks=[_block("Some document text without any tables.")],
        tables=[],  # no tables -> table_index empty -> tables='unavailable'
    )
    manifest = corpus.manifest

    # Verify the DocEntry has tables='unavailable'
    doc_entry = manifest.documents[0]
    assert doc_entry.tables == "unavailable", (
        f"Expected tables='unavailable' for doc with no tables, got {doc_entry.tables!r}"
    )

    ledger = RetrievalLedger()
    faults = detect_structural_inconsistencies(corpus, manifest, ledger)

    assert faults == [], (
        f"D-STR5: doc with unavailable table tier must produce zero faults, got {len(faults)}"
    )


# ---------------------------------------------------------------------------
# Test 4: Single basis cell abstains (no_comparison_basis)
# ---------------------------------------------------------------------------

def test_single_basis_abstains(tmp_path):
    """Table with only ONE non-aggregate data cell -> no Fault (no_comparison_basis).

    Per the algorithm: aggregate detection requires >= 2 independent basis cells
    (after deduplication) to make a meaningful recompute. With only one basis cell,
    the algorithm cannot verify a sum/max/min — it must abstain (return no Fault).
    """
    # Table with only ONE data row + one aggregate row
    # Header: ["Impurity", "% w/w"]
    # Row 1: ["Compound A", "0.10%"]    <- only one basis cell
    # Row 2: ["Total Impurities", "0.20%"]  <- aggregate (only 1 basis -> abstain)
    corpus = _corpus_with_table(
        tmp_path,
        doc_id="doc-single-basis",
        headers=["Impurity", "% w/w"],
        data_rows=[
            ["Compound A", "0.10%"],         # only one non-aggregate row
            ["Total Impurities", "0.20%"],   # aggregate: 1 basis -> abstain
        ],
    )
    manifest = _manifest_for(corpus, "doc-single-basis")
    ledger = RetrievalLedger()

    faults = detect_structural_inconsistencies(corpus, manifest, ledger)

    assert faults == [], (
        f"With only one basis cell, detector must abstain (no_comparison_basis), "
        f"got {len(faults)} fault(s)"
    )


# ---------------------------------------------------------------------------
# Test 5: D-GRD1 fixture cosine regime — fixture realistic text check
# ---------------------------------------------------------------------------

def test_fixture_cosine_regime():
    """D-GRD1 constraint 3: fixture text must be realistic enough for bge-m3 cosine regime.

    Verifies the doc_b.docx synthetic fixture has sufficient domain vocabulary
    to exercise the dense-cosine thresholds. Does NOT require FAISS index.
    """
    import pathlib

    fixture_dir = pathlib.Path("src/evals/dataset/synthetic_fixture")
    assert fixture_dir.exists(), f"fixture dir must exist: {fixture_dir}"
    try:
        import docx as docx_lib
        doc = docx_lib.Document(str(fixture_dir / "doc_b.docx"))
        full_text = " ".join(para.text for para in doc.paragraphs)
        domain_terms = [
            "impurity", "specification", "NMT", "accuracy", "method validation",
            "dissolution", "stability", "validation", "analytical procedure",
        ]
        present = [t for t in domain_terms if t.lower() in full_text.lower()]
        assert len(present) >= 3, (
            f"D-GRD1: fixture doc_b.docx must contain at least 3 domain vocabulary terms "
            f"for non-degenerate bge-m3 cosine regime. Present: {present}"
        )
        assert len(full_text) >= 100, (
            f"D-GRD1: fixture text must be at least 100 chars, got {len(full_text)}"
        )
    except ImportError:
        pytest.skip("python-docx not available for cosine-regime check")
