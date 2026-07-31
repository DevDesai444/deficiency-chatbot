"""Deterministic generator for the planted-deficiency DOCX fixture.

Writes `src/evals/dataset/docs/mini_spec.docx`: a one-page mock impurities
specification that plants exactly 3 hand-verifiable deficiencies, labeled in
`src/evals/dataset/minispec.deficiencies.json` (doc_id="minispec"). This is
Phase 0 eval-harness breadth expansion (see `docs/eval/PITFALLS.md` #11, the
n=1 overfitting trap) -- it establishes the DOCX-format labeled target that
Phase 1's DOCX parse path (INGEST-02) will be gated against.

NOTE: this script does NOT parse DOCX. It only WRITES one, using
`python-docx` (`from docx import Document`). DOCX parsing is Phase 1 scope;
at Phase 0 this document exists in the eval set as a target the harness
records as a parse-fidelity miss until a DOCX parser lands.

Idempotent: re-running overwrites the previous `mini_spec.docx` in place with
the same deterministic paragraph/table content -- no random or time-based
values are written.

Planted deficiencies (kept in lockstep with minispec.deficiencies.json):
  1. "Total impurities" row = 0.14%, while component row "Impurity B" = 0.15%
     -- the total is less than its own largest component, an arithmetic
     impossibility (family = cross_reference_integrity).
  2. "Maximum" row states 0.15%, while the true maximum across the named
     component rows (Impurity C = 0.19%) is higher -- a stated-max mismatch
     (family = cross_reference_integrity).
  3. Narrative claims "Accuracy of the method was established at three
     concentration levels." with no accuracy result table anywhere in the
     document (family = absence_of_evidence).

EMERGENT (not deliberately planted, but factually present and hand-verified):
  4. "Impurity B" row: Result 0.15% > Limit 0.10% -- a direct specification
     exceedance (family = cross_reference_integrity, MS-04). B's 0.15% was
     authored as supporting data for deficiency #1 (the largest single
     component), but 0.15% against its own 0.10% limit is ALSO a real
     deficiency the original 3-item ground truth did not enumerate. Added to
     minispec.deficiencies.json 2026-07-30 after a reviewed live detection run
     caught it (see docs/eval/BASELINE.md) -- a measurement instrument that
     omits a real deficiency the document contains cannot measure "catch all".
     The document BYTES are unchanged; only the ground-truth labeling grew.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document

OUTPUT_PATH = Path(__file__).parent / "dataset" / "docs" / "mini_spec.docx"

_TABLE_HEADER = ("Impurity", "Result (%)", "Limit")

# (label, result, limit) - "Total impurities" and "Maximum" are summary rows
# (not named impurities) but share the same 3-column shape as the rest.
_TABLE_ROWS = (
    ("Impurity A", "0.08", "0.10"),
    ("Impurity B", "0.15", "0.10"),
    ("Impurity C", "0.19", "0.20"),
    ("Total impurities", "0.14", "0.20"),
    ("Maximum", "0.15", "N/A"),
)

_HEADING_TEXT = "Mini Impurities Specification - Mock Drug Substance"
_NARRATIVE_TEXT = (
    "This specification lists acceptance criteria for named and unspecified "
    "impurities observed in a mock drug-substance batch, used to validate the "
    "eval harness's DOCX-format ground truth."
)
# Planted deficiency 3 (absence_of_evidence): this claim is made, but no
# accuracy result table or data exists anywhere in the document.
_ACCURACY_CLAIM_TEXT = "Accuracy of the method was established at three concentration levels."


def build_document() -> Document:
    """Construct the mock impurities-specification Document in memory."""
    document = Document()
    document.add_heading(_HEADING_TEXT, level=1)
    document.add_paragraph(_NARRATIVE_TEXT)
    document.add_paragraph(_ACCURACY_CLAIM_TEXT)

    table = document.add_table(rows=1, cols=len(_TABLE_HEADER))
    table.style = "Table Grid"
    for cell, heading in zip(table.rows[0].cells, _TABLE_HEADER):
        cell.text = heading

    for label, result, limit in _TABLE_ROWS:
        row_cells = table.add_row().cells
        for cell, value in zip(row_cells, (label, result, limit)):
            cell.text = value

    return document


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(str(OUTPUT_PATH))
    print(f"wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
