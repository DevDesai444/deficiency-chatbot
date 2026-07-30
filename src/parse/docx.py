"""DOCX parse path -- converges byte-for-byte on parse/pdf.py's document dict (INGEST-02).

DOCX has no geometry, so reading order == document (XML body) order and bbox/page are null
(D-20 -- never synthetic). This module is the sibling of parse/pdf.py: `ingest/` dispatches to
`parse.pdf.extract_pdf` for PDFs and `parse.docx.extract_docx` for DOCX, and both emit the
identical `{filename, page_count, toc, pages:[{..., blocks, tables, figures}]}` dict.

Merged-cell fidelity (INGEST-05, D-31) uses python-docx's `_tc` grid identity: every (row,col)
a merge spans returns the SAME underlying `<w:tc>` element, so the origin is the first
row-major coordinate that element appears at, and covered coordinates map to it.

SECURITY (T-01-05, XXE): python-docx reads the .docx zip in-memory via lxml and does NOT resolve
external entities or fetch over the network for the document parts read here. The decompression
-bomb / zip-slip guard (`ingest.limits.docx_zip_guard`, Plan 02) runs BEFORE this parser, invoked
by the corpus walker (Plan 09).
"""
from __future__ import annotations

from pathlib import Path

import docx
from docx.table import Table
from docx.text.paragraph import Paragraph

from schemas.documents import ExtractedTable, LayoutBlock
from schemas.llm import ParseFailed


def iter_block_items(document):
    """Yield ('paragraph'|'table', element) in true reading order.

    MUST walk the body: `.paragraphs`/`.tables` are separate lists that lose interleaving
    (a paragraph between two tables cannot be reconstructed from them).
    """
    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            yield ("paragraph", child)
        elif tag == "tbl":
            yield ("table", child)


def _origin_map(tc_grid: list[list]) -> dict[tuple[int, int], tuple[int, int]]:
    """Map every (row,col) to its merge-origin coordinate via `_tc` OBJECT IDENTITY.

    `id(cell._tc)` in a transient loop is UNSAFE -- python-docx hands back short-lived proxies
    that get garbage-collected and whose id() is then reused, producing FALSE merges. `tc_grid`
    holds every `_tc` reference alive, so comparison by `is` is correct.
    """
    origins: list[tuple[object, tuple[int, int]]] = []  # (tc, coord) -- holds refs alive
    coord_to_origin: dict[tuple[int, int], tuple[int, int]] = {}
    for r, row in enumerate(tc_grid):
        for c, tc in enumerate(row):
            found = None
            for utc, coord in origins:
                if utc is tc:
                    found = coord
                    break
            if found is None:
                origins.append((tc, (r, c)))
                coord_to_origin[(r, c)] = (r, c)
            else:
                coord_to_origin[(r, c)] = found
    return coord_to_origin


def _docx_table_to_extracted(tbl_element, order: int, document) -> dict:
    """Reconstruct one `<w:tbl>` into an ExtractedTable dict with merged-origin mapping (D-31).

    Raises on a complex/unreadable merge (python-docx IndexError, issues #1434/#992/#232);
    `extract_docx` catches it and emits a typed ParseFailed marker instead of a corrupt table.
    """
    table = Table(tbl_element, document)
    tc_grid = [[cell._tc for cell in row.cells] for row in table.rows]   # hold refs -> identity-safe
    text_grid = [[cell.text for cell in row.cells] for row in table.rows]
    coord_to_origin = _origin_map(tc_grid)

    n_rows = len(text_grid)
    n_cols = len(text_grid[0]) if text_grid else 0
    merged_origins = {
        f"{r},{c}": [orc[0], orc[1]]
        for (r, c), orc in coord_to_origin.items()
        if orc != (r, c)  # only COVERED (non-origin) coords are keys (D-31)
    }

    def _origin_cells(r: int) -> list[str]:
        return [text_grid[r][c] for c in range(len(text_grid[r])) if coord_to_origin[(r, c)] == (r, c)]

    headers = _origin_cells(0) if n_rows else []
    rows = [_origin_cells(r) for r in range(1, n_rows)]
    return ExtractedTable(
        title="", headers=headers, rows=rows, page=None, kind="grid",
        table_id=f"docx-t{order}", merged_origins=merged_origins,
        bbox=None, n_cols=n_cols, n_rows=n_rows,
    ).model_dump()


def _docx_toc(document) -> list[dict]:
    """Heading-styled paragraphs -> TOC entries (may be empty; flat docs are normal, D-28/D-30)."""
    toc: list[dict] = []
    for para in document.paragraphs:
        name = (para.style.name if para.style else "") or ""
        if name.startswith("Heading") and para.text.strip():
            try:
                level = int(name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            toc.append({"level": level, "title": para.text.strip(), "page": None})
    return toc


def extract_docx(path: str | Path) -> dict:
    """Parse a DOCX into the SAME structured-document dict `parse.pdf.extract_pdf` emits (D-20)."""
    document = docx.Document(str(path))
    blocks: list[dict] = []
    tables: list[dict] = []
    order = 0
    for kind, el in iter_block_items(document):
        if kind == "paragraph":
            text = Paragraph(el, document).text or ""
            if text.strip():
                blocks.append(
                    LayoutBlock(role="paragraph", text=text, bbox=None, page=None,
                               reading_order=order, lines=[]).model_dump()
                )
                order += 1
        else:  # table
            try:
                tables.append(_docx_table_to_extracted(el, order, document))
            except Exception as exc:  # noqa: BLE001 -- Pitfall 5: never crash; mark parsed_partial (D-17)
                tables.append(
                    {"_parse_failed": ParseFailed(
                        layer="parse.docx",
                        reason=f"table {order}: complex merge unreadable: {exc}",
                        raw_output="",
                    ).model_dump()}
                )
            order += 1
    return {
        "filename": Path(path).name,
        "page_count": None,
        "toc": _docx_toc(document),
        "pages": [
            {
                "page_number": None,
                "page_label": "",
                "width": None,
                "height": None,
                "rotation": 0,
                "source": "python-docx",
                "is_scanned": False,
                "blocks": blocks,
                "tables": tables,
                "figures": [],
            }
        ],
    }
